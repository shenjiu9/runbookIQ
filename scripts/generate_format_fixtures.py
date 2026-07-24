from docx import Document
from PIL import Image, ImageDraw, ImageFont


def generate_ocr_fixture() -> None:
    image = Image.new("RGB", (1200, 260), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default(size=72)
    draw.text((40, 70), "OCR TOKEN VISION-5832", fill="black", font=font)
    image.save("/tmp/runbookiq-ocr-proof.png")


def generate_docx_fixture() -> None:
    document = Document()
    document.add_heading("DOCX Deployment Proof", level=1)
    document.add_paragraph("Unique document token DOCX-9146 confirms Word parsing works.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Check"
    table.cell(0, 1).text = "Result"
    table.cell(1, 0).text = "Tokyo"
    table.cell(1, 1).text = "Healthy"
    document.save("/tmp/runbookiq-docx-proof.docx")


if __name__ == "__main__":
    generate_ocr_fixture()
    generate_docx_fixture()
