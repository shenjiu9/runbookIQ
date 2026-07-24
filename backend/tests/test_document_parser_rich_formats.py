import io

from docx import Document
from PIL import Image
from pypdf import PdfWriter

from runbookiq.ingestion.parser import DocumentParser


class FixedOcr:
    def __init__(self) -> None:
        self.image_calls = 0
        self.pdf_calls = 0

    def extract_image(self, content: bytes) -> str:
        self.image_calls += 1
        return "图片中的报销上限是 600 元。"

    def extract_pdf(self, content: bytes) -> str:
        self.pdf_calls += 1
        return "扫描版手册要求先检查上一轮容器日志。"


def docx_fixture() -> bytes:
    document = Document()
    document.add_heading("员工差旅制度", level=1)
    document.add_paragraph("住宿费用按城市等级报销。")
    document.add_heading("费用标准", level=2)
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "城市"
    table.cell(0, 1).text = "上限"
    table.cell(1, 0).text = "上海"
    table.cell(1, 1).text = "600 元"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def image_fixture() -> bytes:
    buffer = io.BytesIO()
    Image.new("RGB", (20, 20), color="white").save(buffer, format="PNG")
    return buffer.getvalue()


def scanned_pdf_fixture() -> bytes:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def test_docx_preserves_headings_paragraphs_and_table_text() -> None:
    sections = DocumentParser().parse(
        filename="travel.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        content=docx_fixture(),
    )

    assert [section.section_path for section in sections] == [
        "员工差旅制度",
        "员工差旅制度 / 费用标准",
    ]
    assert "住宿费用按城市等级报销" in sections[0].text
    assert "上海 | 600 元" in sections[1].text


def test_image_is_converted_to_searchable_text_through_ocr() -> None:
    ocr = FixedOcr()
    sections = DocumentParser(ocr_engine=ocr).parse(
        filename="expense.png",
        content_type="image/png",
        content=image_fixture(),
    )

    assert ocr.image_calls == 1
    assert sections[0].text == "图片中的报销上限是 600 元。"


def test_pdf_without_text_layer_falls_back_to_ocr() -> None:
    ocr = FixedOcr()
    sections = DocumentParser(ocr_engine=ocr).parse(
        filename="scanned.pdf",
        content_type="application/pdf",
        content=scanned_pdf_fixture(),
    )

    assert ocr.pdf_calls == 1
    assert sections[0].text == "扫描版手册要求先检查上一轮容器日志。"
